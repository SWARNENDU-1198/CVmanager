import os
from pypdf import PdfReader
from docx import Document

def clean_text(text: str) -> str:
    if not text:
        return ""
    # Normalize whitespaces
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        cleaned_line = " ".join(line.split())
        cleaned_lines.append(cleaned_line)
    
    # Reconstruct paragraphs but remove blank lines overflow
    content = "\n".join(cleaned_lines)
    while "\n\n\n" in content:
        content = content.replace("\n\n\n", "\n\n")
    return content.strip()

def extract_text_from_pdf(filepath: str) -> str:
    reader = PdfReader(filepath)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(filepath: str) -> str:
    doc = Document(filepath)
    text_list = []
    for para in doc.paragraphs:
        if para.text:
            text_list.append(para.text)
    
    # Extract text from tables as well
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text_list.append(" | ".join(row_text))
                
    return "\n".join(text_list)

def extract_text_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_image(filepath: str) -> str:
    import base64
    from langchain_core.messages import HumanMessage
    from backend.ai.llm_factory import get_llm, invoke_llm_with_retry
    
    ext = os.path.splitext(filepath)[1].lower()
    mime_type = "image/png" if ext == ".png" else "image/jpeg"
    
    with open(filepath, "rb") as image_file:
        base64_image = base64.b64encode(image_file.read()).decode("utf-8")
        
    message = HumanMessage(
        content=[
            {"type": "text", "text": "Extract all text from this resume image. Do not summarize or format, just output the raw text exactly as printed in the image."},
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{base64_image}"}
            }
        ]
    )
    
    llm = get_llm(temperature=0.0, model_type="vision")
    response = invoke_llm_with_retry(llm, [message])
    
    content = response.content
    if isinstance(content, list):
        text_parts = []
        for part in content:
            if isinstance(part, str):
                text_parts.append(part)
            elif isinstance(part, dict) and "text" in part:
                text_parts.append(part["text"])
        return "\n".join(text_parts)
    return str(content)



def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    raw_text = ""
    
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(filepath)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(filepath)
    elif ext == ".txt":
        raw_text = extract_text_from_txt(filepath)
    elif ext in [".png", ".jpg", ".jpeg"]:
        raw_text = extract_text_from_image(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    return clean_text(raw_text)
